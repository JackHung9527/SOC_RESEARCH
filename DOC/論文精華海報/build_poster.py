# -*- coding: utf-8 -*-
"""產生 A4 直式「論文精華海報」DOCX（繁體中文、單頁）。

版面：A4 直式、窄邊界；上方跨頁寬標題區，中段左右雙欄（以 1x2 表格佈局），
下方跨頁寬放三法軌跡圖、決策表與結論。所有數據均取自論文六章之實測結果
（Rounds 40／41 板端整輪實測、rounds 1–41 跨輪紀錄）。

輸出：論文精華海報_A4.docx（同目錄）
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
OUT = os.path.join(HERE, "論文精華海報_A4.docx")

CJK = "Microsoft JhengHei"
LAT = "Segoe UI"

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x5C, 0x9A)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ---- 版面尺寸與字級（單頁調校用旋鈕）----
MARGIN = dict(top=0.7, bottom=0.6, left=1.0, right=1.0)     # cm
FULL_W = Cm(19.0)
COL_W = Cm(9.25)
FS_HEAD = 8.8        # 區塊標題
FS_BODY = 7.4        # 內文
FS_TBL = 6.8         # 表格內文
FS_NOTE = 6.3        # 表註／圖說
W_FIG1 = 6.7         # 圖 1 寬 (cm)
W_FIG2 = 7.8         # 圖 2 寬 (cm)
W_FIG3 = 13.5        # 圖 3 寬 (cm)


# ------------------------------------------------------------------ helpers
def set_font(run, size, bold=False, color=BLACK, italic=False, latin=LAT):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:eastAsia"), CJK)


def fmt_par(par, space_before=0, space_after=1.0, line=1.0, align=None,
            keep_next=False):
    pf = par.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if align is not None:
        par.alignment = align
    pf.keep_with_next = keep_next
    return par


def shade(element, hexcolor):
    """對段落或儲存格套底色。"""
    if hasattr(element, "_p"):
        pr = element._p.get_or_add_pPr()
    else:
        pr = element._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    pr.append(shd)


def cell_margins(cell, top=30, bottom=30, left=70, right=70):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left),
                     ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcpr.append(mar)


def table_borders(table, color="BFBFBF", sz=4, none=False):
    tblpr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:val"), "none" if none else "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblpr.append(borders)


def no_split(table):
    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        trpr.append(OxmlElement("w:cantSplit"))


def spacer(container, size=1):
    """表格之間的極小空段落，避免相鄰表格被 Word 併成同一表格。"""
    par = container.add_paragraph()
    fmt_par(par, space_before=0, space_after=0, line=1.0)
    set_font(par.add_run(""), size)
    return par


def section_head(container, text, size=FS_HEAD):
    par = container.add_paragraph()
    fmt_par(par, space_before=1.6, space_after=1.4, keep_next=True)
    run = par.add_run("  " + text)
    set_font(run, size, bold=True, color=WHITE)
    shade(par, "2E5C9A")
    return par


def body(container, text, size=FS_BODY, bullet=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1.2, color=BLACK,
         bold_head=None):
    par = container.add_paragraph()
    fmt_par(par, space_after=space_after, align=align, line=1.0)
    if bullet:
        par.paragraph_format.left_indent = Cm(0.32)
        par.paragraph_format.first_line_indent = Cm(-0.32)
        run = par.add_run("▪ ")
        set_font(run, size, bold=True, color=BLUE)
    if bold_head:
        set_font(par.add_run(bold_head), size, bold=True, color=NAVY)
    set_font(par.add_run(text), size, color=color)
    return par


def caption(container, text, size=FS_NOTE):
    par = container.add_paragraph()
    fmt_par(par, space_before=0.5, space_after=1.5,
            align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0)
    set_font(par.add_run(text), size, color=GREY)
    return par


def picture(container, path, width_cm):
    par = container.add_paragraph()
    fmt_par(par, space_before=1, space_after=0,
            align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
    par.add_run().add_picture(path, width=Cm(width_cm))
    return par


def data_table(container, rows, widths_cm, size=FS_TBL, head_fill="D6E1F1",
               aligns=None):
    """rows[0] 為表頭；欄位字串前綴 '*' 代表該格加粗強調。"""
    tbl = container.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    table_borders(tbl)
    amap = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER,
            "r": WD_ALIGN_PARAGRAPH.RIGHT}
    aligns = aligns or ["l"] + ["c"] * (len(rows[0]) - 1)
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = tbl.cell(i, j)
            cell.width = Cm(widths_cm[j])
            cell_margins(cell, top=12, bottom=12, left=55, right=55)
            par = cell.paragraphs[0]
            fmt_par(par, space_after=0, line=1.0,
                    align=amap[aligns[j] if i else "c"])
            emph = txt.startswith("*")
            run = par.add_run(txt[1:] if emph else txt)
            set_font(run, size, bold=(i == 0 or emph),
                     color=NAVY if i == 0 else BLACK)
            if i == 0:
                shade(cell, head_fill)
    no_split(tbl)
    return tbl


def note(container, text, size=FS_NOTE):
    par = container.add_paragraph()
    fmt_par(par, space_before=0.8, space_after=1.2,
            align=WD_ALIGN_PARAGRAPH.LEFT, line=1.0)
    set_font(par.add_run(text), size, color=GREY)
    return par


def banner(container, fill, runs, width, pad=55, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """整寬色塊（標題亮點條／結論框）。runs 為 (文字, 粗體, 字級, 顏色) 串列。"""
    tbl = container.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = width
    cell = tbl.cell(0, 0)
    cell.width = width
    table_borders(tbl, none=True)
    shade(cell, fill)
    cell_margins(cell, top=pad, bottom=pad, left=120, right=120)
    par = cell.paragraphs[0]
    fmt_par(par, space_after=0, align=align, line=1.04)
    for txt, bold, size, color in runs:
        set_font(par.add_run(txt), size, bold=bold, color=color)
    no_split(tbl)
    return tbl


# --------------------------------------------------------------------- build
doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin, sec.bottom_margin = Cm(MARGIN["top"]), Cm(MARGIN["bottom"])
sec.left_margin, sec.right_margin = Cm(MARGIN["left"]), Cm(MARGIN["right"])

style = doc.styles["Normal"]
style.font.name = LAT
style.font.size = Pt(FS_BODY)
style.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)

# ---- 標題區 ----
p = doc.add_paragraph()
fmt_par(p, space_after=0.5, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
set_font(p.add_run("鋰電池 SOC 估測方法之比較與嵌入式實作"), 17, bold=True,
         color=NAVY)

p = doc.add_paragraph()
fmt_par(p, space_after=1, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
set_font(p.add_run("A Comparative Study and Embedded Implementation of "
                   "SOC Estimation Methods for Lithium-ion Batteries"),
         8.2, italic=True, color=BLUE, latin="Times New Roman")

p = doc.add_paragraph()
fmt_par(p, space_after=1.8, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
set_font(p.add_run("中原大學 資訊工程學系碩士班　　研究生：洪大甲　　"
                   "指導教授：鄭維凱"), 8.4)

banner(doc, "1F3864", [
    ("同電池、同協定、同微控制器的公平比較：", True, 8.2, WHITE),
    ("EKF 全倍率 RMSE 0.22–0.90%、庫倫 1.7–2.0%、動態阻抗線上不可行"
     "──量測域一致性凌駕演算法選擇。", False, 8.2, WHITE),
], FULL_W, pad=45, align=WD_ALIGN_PARAGRAPH.CENTER)
spacer(doc)

# ---- 雙欄主體 ----
main = doc.add_table(rows=1, cols=2)
main.autofit = False
table_borders(main, none=True)
left, right = main.cell(0, 0), main.cell(0, 1)
left.width = right.width = COL_W
cell_margins(left, top=0, bottom=0, left=0, right=120)
cell_margins(right, top=0, bottom=0, left=120, right=0)
for c in (left, right):
    c.paragraphs[0]._p.getparent().remove(c.paragraphs[0]._p)

# ======================= 左欄 =======================
section_head(left, "1　研究動機與缺口")
body(left, "現有研究多於 PC 端、理想量測條件下評比 SOC 演算法，且各研究的電池、協定"
           "與硬體皆不相同，因而留下兩個缺口：")
body(left, "演算法在真實部署硬體上的表現未被驗證。", bullet=True,
     bold_head="評估環境失真──")
body(left, "精度差異與電池、協定、硬體差異混雜，無法歸因於方法本身。",
     bullet=True, bold_head="缺乏公平基準──")

section_head(left, "2　自動化測試平台與量測基礎")
picture(left, os.path.join(FIG, "poster_fig1_arch.png"), W_FIG1)
caption(left, "圖 1　自動化跨輪電池測試平台架構")
body(left, "INA226 前端 14 點線性內插校正，全範圍電流誤差 < 1‰。",
     bullet=True, bold_head="量測校正：")
body(left, "CC-CV 充電 → 靜置 → 0.5–2.0C 四倍率放電，一輪 16–18 h 無人值守；"
           "放電中每 60 s 施加 dV/dI 擾動，供動態阻抗建表。",
     bullet=True, bold_head="跨輪協定：")
body(left, "三輪 CV < 0.52%（噪聲下界）；41 輪老化 4.6%，為噪聲的 9 倍。",
     bullet=True, bold_head="平台再現性：")

section_head(left, "3　三方法之嵌入式實作（STM32G071 @64 MHz）")
body(left, "int64 電荷累加、純整數運算，具充飽自動重錨；兼作板端真值。",
     bullet=True, bold_head="庫倫計數：")
body(left, "一階 RC、狀態二維、增益免矩陣求逆；GITT 21 點 OCV 表為觀測方程。",
     bullet=True, bold_head="EKF：")
body(left, "複用既有 dV/dI 擾動，二次式 Z–SOC 反解，無須靜置與初值。",
     bullet=True, bold_head="動態阻抗：")
body(left, "共用前端、分離估測、統一輸出；三法共讀同一筆量測。",
     bullet=True, bold_head="並行架構：")

# ======================= 右欄 =======================
section_head(right, "4　強健性壓力測試（板端／PC 重放實測）")
data_table(right, [
    ["壓力測試", "庫倫計數", "EKF", "動態阻抗"],
    ["初值錯誤恢復", "不收斂", "*單步拉回", "首事件 ≤60 s"],
    ["C-rate 切換尖峰", "*≤0.21%", "≤0.48%", "10–48%"],
    ["噪聲下 SOC 抖動", "*≤0.054%", "≤0.27%", "約 53%"],
], widths_cm=[2.75, 2.05, 1.9, 2.15], aligns=["l", "c", "c", "c"])
note(right, "表 1　庫倫無修正機制，故加入充飽自動重錨（整輪 4/4 次正確觸發）；EKF 以"
            "共變異數濾波使兩類誤差皆有界；動態阻抗於切換與噪聲下皆因二次反解的分枝"
            "翻轉而失控。")
spacer(right)

section_head(right, "5　核心發現：量測域一致性凌駕演算法選擇")
picture(right, os.path.join(FIG, "poster_fig2_zdomain.png"), W_FIG2)
caption(right, "圖 2　動態阻抗之量測域對照（Round 40 實測）")
body(right, "參數辨識域（台架負載端）與部署量測域（板端電池端子）之間約 24 mΩ 的"
            "線材／接點電阻，使 EKF 產生隨倍率放大的系統性偏差，並使動態阻抗對照表"
            "整組失效；兩域擬合僅常數項平移──曲線形狀由電芯決定，常數項由量測鏈"
            "決定。")
body(right, "R₀ 換算至部署域後，EKF 的 RMSE 由最高 4.48% 降至全倍率 < 1%，"
            "倍率相依偏差完全消失；「離線建表、線上查表」的跨量測鏈搬移，"
            "是嵌入式 SOC 估測的隱形失效模式。", bold_head="獨立複測證實：")

# ---- 跨欄底部 ----
spacer(doc)
section_head(doc, "6　三法 SOC 估測軌跡實測對照")
picture(doc, os.path.join(FIG, "poster_fig3_soc.png"), W_FIG3)
caption(doc, "圖 3　三法每秒輸出與台架庫倫真值之對照（Round 41 板端實測）：EKF 與真值"
             "幾乎重合，庫倫呈刻度差造成的輕微負偏，動態阻抗則出現鏡像跳變與夾限平台。")

section_head(doc, "7　精度─資源權衡與方法選用決策（同硬體實測，本研究核心產出）")
data_table(doc, [
    ["方法", "RMSE (%)", "最大誤差 (%)", "Flash (B)", "RAM (B)",
     "cycles／次更新", "適用情境與主要代價（實測依據）"],
    ["庫倫計數", "1.7–2.0", "3.0–3.5", "+1,892", "+24", "305（4.8 µs）",
     "極低成本／算力（8-bit、Flash < 8 KB）：資源下界；無自我修正，精度上限"
     "由電流量測鏈刻度（±2%）決定"],
    ["*EKF（一階 RC）", "*0.22–0.90", "*0.6–1.9", "+2,336", "+40",
     "15,187（237 µs）",
     "中階且需自初值錯誤恢復：全倍率精度最佳、錯誤初值單步拉回；須 GITT 建 "
     "OCV 表，參數必須與部署量測鏈同域"],
    ["動態阻抗", "23–31（線上）", "約 96", "+1,540", "+48",
     "442／6,525（有事件）",
     "無靜置機會、實驗成本敏感：無須初值、首事件 ≤60 s 重估；線上單獨反推對平坦"
     "曲線電芯不可行，宜作低 SOC 端粗校正與老化指標"],
    ["三法並行", "—", "—", "+4,936", "+112", "約 22,000",
     "高可靠度、資源充裕：三法互補，峰值僅佔每秒運算預算 0.034%（餘裕逾三個"
     "數量級）；整合複雜度最高"],
], widths_cm=[2.15, 1.95, 1.6, 1.5, 1.25, 2.35, 8.2],
    aligns=["l", "c", "c", "r", "r", "c", "l"], size=6.4)
note(doc, "表 2　精度為 Round 41 板端整輪四倍率之範圍（庫倫的恆負偏源自兩條電流量測鏈約 "
          "+1.2% 的刻度差，非演算法誤差；動態阻抗離線、假設分枝正確之理想上界為 6–10%）；"
          "資源為對共用韌體骨架之差分量測，運算量取 Rounds 40–41 共 159,300 次更新之中位數。")
spacer(doc)

banner(doc, "EAF0F8", [
    ("結論　", True, FS_BODY, NAVY),
    ("不存在單一最優的 SOC 估測方法，最優選擇取決於硬體資源與應用對初值恢復、靜置"
     "機會與實驗成本的要求；本研究並指出演算法的精度上限往往不由演算法決定──量測鏈"
     "與辨識域／部署域的一致性，對誤差的貢獻數倍於方法之間的差距。　未來工作：SOH "
     "估測延伸、LFP 遲滯建模、輕量資料驅動融合。", False, FS_BODY, BLACK),
], FULL_W, pad=45)

# Word 於文件結尾（表格之後）必有一個段落；壓成 1 pt 精確行高避免溢出到第二頁
tail = doc.add_paragraph()
fmt_par(tail, space_before=0, space_after=0)
tail.paragraph_format.line_spacing = Pt(1)
set_font(tail.add_run(""), 1)

doc.save(OUT)
print(f"[done] {OUT}")
