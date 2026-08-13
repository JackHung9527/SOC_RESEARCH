# -*- coding: utf-8 -*-
"""
Build the TANET conference paper (version 2026081002) directly ON TOP OF the official
TANET_Format_Paper.docx template found in refer/.

Differences vs. build_tanet.py (the 20260722 build):
  * The Chinese abstract (摘要 + 關鍵詞) is removed - the TANET template carries an
    English Abstract only.
  * Every paragraph geometry (line height, spacing, indent, font size, column layout,
    margins) now comes from the template's own named styles instead of hard-coded
    values, so the output matches TANET_Format_Paper.docx exactly:
        paper title / 內文1 / Affiliation / 標題 51 / Abstract / key words
        標題 11 (auto "I.") / 標題 21 (auto "A.") / 本文1 (10pt, line 228 auto)
        equation / table head (auto "TABLE I.") / table col head / table copy
        figure caption (auto "Figure 1.") / references (auto "[1]")
  * Section / table / figure / reference numbers are Word auto-numbering fields from
    the template's numbering.xml, so they renumber themselves in Word.
"""
import os
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
TEMPLATE = os.path.join(HERE, "refer", "TANET_Format_Paper.docx")
OUT_DOCX = os.path.join(
    HERE,
    "A_Comparative_Study_and_Embedded_Implementation_of_SOC_Estimation_Methods_"
    "for_Lithium-ion_Batteries_2026081002.docx")

TIMES = "Times New Roman"

# style names as they appear in the template
S_TITLE = "paper title"
S_CENTER = "內文1"
S_AFFIL = "Affiliation"
S_HEADX = "標題 51"      # unnumbered heading (Abstract / References)
S_ABS = "Abstract"
S_KEYS = "key words"
S_H1 = "標題 11"        # auto "I."  centered small-caps
S_H2 = "標題 21"        # auto "A."  left italic
S_BODY = "本文1"
S_EQ = "equation"
S_TBLHEAD = "table head"        # auto "TABLE I."
S_TBLCOLHEAD = "table col head"
S_TBLCOPY = "table copy"
S_FIGCAP = "figure caption"     # auto "Figure 1."
S_REF = "references"            # auto "[1]"


# ------------------------------------------------------------------ template prep
doc = Document(TEMPLATE)
body = doc.element.body

# locate the two paragraphs that carry a sectPr (title-block end, body end)
sect_paras = [el for el in body.findall(qn('w:p'))
              if el.find(qn('w:pPr')) is not None
              and el.find(qn('w:pPr')).find(qn('w:sectPr')) is not None]
assert len(sect_paras) == 2, f"expected 2 in-body sectPr paragraphs, got {len(sect_paras)}"
ANCHOR_TITLE, ANCHOR_BODY = sect_paras          # insert before these

# keep the trailing paragraph of the final (balancing) section
TAIL = body.findall(qn('w:p'))[-1]
KEEP = {id(ANCHOR_TITLE), id(ANCHOR_BODY), id(TAIL)}

# reference 2-column sectPr (from the body-end paragraph) for the spanning figure
SECT_2COL = deepcopy(ANCHOR_BODY.find(qn('w:pPr')).find(qn('w:sectPr')))

# drop the template's sample picture relationship, then wipe all sample content
for blip in body.iter(qn('a:blip')):
    rid = blip.get(qn('r:embed'))
    if rid:
        try:
            doc.part.drop_rel(rid)
        except Exception:
            pass
for el in list(body):
    if el.tag in (qn('w:p'), qn('w:tbl')) and id(el) not in KEEP:
        body.remove(el)


# ------------------------------------------------------------------- insert helper
_cursor = {"anchor": ANCHOR_TITLE}

# Tracks the most recently emitted paragraph so a spanning figure can hang its
# 2-column sectPr on it instead of creating a throwaway spacer paragraph.
# Reset to None whenever the last emitted element is NOT a paragraph (e.g. a table):
# a sectPr must sit on the last paragraph OF the section, and hanging it on a table's
# caption would wrongly push the table itself into the following section.
_last = {"par": None}


def _at(anchor):
    _cursor["anchor"] = anchor


def _runs_from(text):
    """Split **bold** spans -> list of (chunk, is_bold)."""
    out = []
    for i, chunk in enumerate(text.split("**")):
        if chunk:
            out.append((chunk, i % 2 == 1))
    return out


def p(style, text="", *, align=None, size=None, bold=None, italic=None,
      keep_next=False):
    """Create a paragraph with a template style and move it in front of the cursor."""
    par = doc.add_paragraph()               # lands just before the body sectPr
    _cursor["anchor"].addprevious(par._p)   # relocate to the wanted spot
    par.style = style
    if align is not None:
        par.alignment = align
    if keep_next:
        par.paragraph_format.keep_with_next = True
    for chunk, b in _runs_from(text):
        r = par.add_run(chunk)
        if size is not None:
            r.font.size = Pt(size)
        if bold is not None or b:
            r.font.bold = True if b else bold
        if italic is not None:
            r.font.italic = italic
    _last["par"] = par
    return par


def end_section_at(par, ncols):
    """Attach a CONTINUOUS sectPr to `par`; the section it ends uses `ncols` columns."""
    sect = deepcopy(SECT_2COL)
    t = sect.find(qn('w:type'))
    if t is None:
        t = OxmlElement('w:type')
        sect.insert(0, t)
    t.set(qn('w:val'), 'continuous')
    if ncols == 1:
        old = sect.find(qn('w:cols'))
        if old is not None:
            sect.remove(old)
        cols = OxmlElement('w:cols')
        cols.set(qn('w:space'), '720')
        sect.append(cols)
    par._p.get_or_add_pPr().append(sect)


# ------------------------------------------------------------------------- tables
def _border(tc):
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '2')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), '000000')
        b.append(e)
    tcPr.append(b)
    mar = OxmlElement('w:tcMar')
    for side, w in (('top', '0'), ('left', '54'), ('bottom', '0'), ('right', '54')):
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), w)
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tcPr.append(mar)
    va = OxmlElement('w:vAlign')
    va.set(qn('w:val'), 'center')
    tcPr.append(va)


def add_table(caption, rows, widths_twips):
    """caption (auto-numbered TABLE n) above a fixed-layout, centred table."""
    p(S_TBLHEAD, caption, keep_next=True)

    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    _cursor["anchor"].addprevious(t._tbl)

    tblPr = t._tbl.tblPr
    for tag in ('w:tblW', 'w:jc', 'w:tblLayout'):
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(sum(widths_twips)))
    w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)

    grid = t._tbl.find(qn('w:tblGrid'))
    for gc, tw in zip(grid.findall(qn('w:gridCol')), widths_twips):
        gc.set(qn('w:w'), str(tw))

    for ri, row in enumerate(rows):
        trPr = t.rows[ri]._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        if ri == 0:
            trPr.append(OxmlElement('w:tblHeader'))
        rjc = OxmlElement('w:jc')
        rjc.set(qn('w:val'), 'center')
        trPr.append(rjc)
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.width = Pt(widths_twips[ci] / 20.0)
            _border(cell._tc)
            par = cell.paragraphs[0]
            par.style = S_TBLCOLHEAD if ri == 0 else S_TBLCOPY
            par.alignment = (WD_ALIGN_PARAGRAPH.CENTER if (ri == 0 or ci > 0)
                             else WD_ALIGN_PARAGRAPH.LEFT)
            par.paragraph_format.keep_with_next = (ri < len(rows) - 1)
            for chunk, b in _runs_from(val):
                r = par.add_run(chunk)
                if b:
                    r.font.bold = True
    _last["par"] = None          # last emitted element is a table, not a paragraph
    return t


def add_figure(path, width_cm, caption, *, span=False):
    """Insert a picture + auto-numbered caption. span=True -> full page width."""
    if span:
        # Close the running 2-column section. Prefer hanging the sectPr on the
        # preceding body paragraph so no empty spacer line is left behind; fall
        # back to a spacer only when the previous element was a table.
        if _last["par"] is not None:
            end_section_at(_last["par"], 2)
        else:
            sp = p(S_CENTER, "")
            sp.paragraph_format.space_after = Pt(0)
            sp.paragraph_format.line_spacing = Pt(4)
            for r in sp.runs:
                r.font.size = Pt(2)
            end_section_at(sp, 2)

    pic = p(S_CENTER, "", align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True)
    pic.add_run().add_picture(path, width=Cm(width_cm))
    cap = p(S_FIGCAP, caption)

    if span:
        end_section_at(cap, 1)


# ================================================================== TITLE BLOCK
_at(ANCHOR_TITLE)

p(S_TITLE, "A Comparative Study and Embedded Implementation of "
           "SOC Estimation Methods for Lithium-ion Batteries", size=16, bold=True)
p(S_CENTER, "")
p(S_CENTER, "Ta-Chia Hung, Wei-Kai Cheng", size=12)
p(S_AFFIL, "Department of Information and Computer Engineering", size=12)
p(S_AFFIL, "Chung Yuan Christian University, Taoyuan City, Taiwan", size=12)
p(S_CENTER, "g11277602@cycu.edu.tw, wkcheng@cycu.edu.tw", size=12)
p(S_AFFIL, "")
p(S_CENTER, "")


# ======================================================================== BODY
_at(ANCHOR_BODY)

# ------------------------------------------------------------------- Abstract
p(S_HEADX, "Abstract", size=12, bold=True)
p(S_ABS,
  "State of Charge (SOC) estimation is a core function of a battery management "
  "system, yet existing method comparisons are mostly performed offline using a "
  "single accuracy metric and rarely quantify, on the resource-constrained "
  "microcontroller where the algorithm is actually deployed, both accuracy and "
  "resource cost; differing cells, protocols and hardware further prevent a fair "
  "baseline. This work builds an automated multi-round test platform that keeps the "
  "cell, the test protocol and the microcontroller identical across methods, and "
  "implements Coulomb counting, an Extended Kalman Filter (EKF) and a "
  "dynamic-impedance method on the same STM32 under a unified firmware "
  "specification. Full-round on-board measurements show that the EKF attains a "
  "root-mean-square error below 1% (0.22-0.90%) at all four discharge rates - the "
  "best of the three - while Coulomb counting yields 1.7-2.0% and impedance-only "
  "online inversion is ill-posed (23-31%) for this flat Z-SOC cell. The EKF costs "
  "about 50x the compute of Coulomb counting yet consumes only 0.024% of the "
  "1-second update budget at 64 MHz. We further show that measurement-domain "
  "consistency dominates algorithm choice: parameters and look-up tables built in a "
  "different measurement domain than deployment cause systematic error regardless "
  "of algorithm correctness.")
p(S_KEYS, "Keywords: lithium-ion battery, state of charge estimation, Coulomb "
          "counting, extended Kalman filter, dynamic impedance, embedded systems")

# --------------------------------------------------------------- I. Introduction
p(S_H1, "Introduction")
p(S_BODY,
  "Driven by decarbonization and the electrification of transport, lithium-ion "
  "batteries have become the dominant energy-storage element in electric vehicles, "
  "portable electronics and grid storage. The battery management system (BMS) is "
  "therefore indispensable, and among its functions the State of Charge (SOC) - the "
  "ratio of remaining to rated capacity - plays the role of a fuel gauge for range "
  "prediction, energy scheduling and over-charge/over-discharge protection. SOC is "
  "an internal state that cannot be sensed directly; it can only be inferred from "
  "terminal voltage, current and temperature of a highly nonlinear, time-varying "
  "electrochemical cell [1], [4].")
p(S_BODY,
  "Although comparisons of SOC estimation methods are abundant, two gaps remain "
  "insufficiently addressed. The first is evaluation-environment distortion. Most "
  "comparisons are conducted on a PC or in MATLAB and report the root-mean-square "
  "error (RMSE) as almost the only metric, whereas the real deployment target is a "
  "resource-constrained microcontroller unit (MCU). How much Flash and RAM a method "
  "occupies after porting, how many CPU cycles each update costs, and how much "
  "accuracy is lost after fixed-point implementation are seldom quantified. "
  "Notably, commercial BMS ICs generally adopt Coulomb counting with OCV correction "
  "rather than the theoretically more accurate Extended Kalman Filter (EKF) [5] - "
  "an engineering trade-off among memory, computation and cost that is rarely "
  "presented with measured data.")
p(S_BODY,
  "The second gap is unfair benchmarking. Existing cross-method comparisons often "
  "use different cells, protocols and hardware, so the reported accuracies share no "
  "common baseline. The dynamic-impedance method [1], for instance, claims to need "
  "no initial value and to suit embedded use, yet it has never been placed on the "
  "same MCU, the same cell and the same protocol as an EKF for an equal "
  "resource-versus-accuracy comparison.")
p(S_BODY,
  "To address both gaps, this work builds a “same cell, same protocol, same "
  "MCU” fair-comparison platform and quantifies the accuracy-cost trade-off of "
  "embedded SOC estimation. Its contributions are: (1) an automated, long-term "
  "unattended cross-round test platform integrating an STM32, a programmable DC "
  "supply, a programmable electronic load and a high-accuracy current/voltage "
  "sensor with an automation scheduler; (2) an implementation of Coulomb counting, "
  "the EKF and the dynamic-impedance method on one STM32 under a unified firmware "
  "specification; and (3) an equal, reproducible comparison of accuracy, robustness "
  "and embedded resource cost, from which we derive both a method-selection guide "
  "and a key finding that measurement-domain consistency dominates algorithm "
  "choice.")

# ------------------------------------------------------------------ II. Background
p(S_H1, "Background and Related Work")
p(S_H2, "Equivalent-circuit model")
p(S_BODY,
  "A first-order RC (Thevenin) equivalent-circuit model is adopted as the cell "
  "model: an open-circuit-voltage source V_OC(SOC) in series with an ohmic "
  "resistance R0 and one parallel RC branch (R1, C1) that captures polarization. "
  "The terminal voltage is given by (1), where V1 is the RC-branch voltage. This "
  "structure balances fidelity and identifiability for embedded use [6], and the "
  "nonlinear V_OC(SOC) relation is obtained by a pseudo-OCV table [5].")
p(S_H2, "SOC estimation methods")
p(S_BODY,
  "Coulomb counting integrates current to accumulate charge; it is simple and "
  "model-free but has no self-correction and accumulates drift [4]. OCV look-up "
  "maps a rested open-circuit voltage to SOC but requires long relaxation. "
  "Model-based methods, dominated by the Kalman-filter family (KF/EKF/UKF), fuse "
  "current integration with a voltage measurement to correct the estimate and can "
  "recover from a wrong initial value [2], [3], [7], [9]. The dynamic-impedance "
  "method [1] estimates SOC from the instantaneous voltage-to-current ratio during "
  "operation, avoiding both an initial value and long rest, and online impedance "
  "has also been used for low-complexity modeling and aging indication [8]. "
  "Commercial ICs nonetheless favor Coulomb counting with OCV correction for "
  "resource reasons [5], a trade-off this work quantifies on real hardware.")

# ------------------------------------------------------------------- III. Platform
p(S_H1, "Test Platform and Estimator Implementation")
p(S_H2, "System architecture")
p(S_BODY,
  "The platform plays two roles simultaneously: an excitation-and-ground-truth rig "
  "and an embedded estimation target. A programmable DC supply (ITECH IT6302, "
  "CC-CV) and electronic load (ITECH IT8512A+, CC) apply controlled "
  "charge/discharge to a single device-under-test (DUT), while an STM32G071 with an "
  "INA226 current/voltage sensor runs the estimators. Both roles share the same "
  "cell and the same high-power loop, so the instrument ground truth and the MCU "
  "estimate can be aligned point-by-point. Sensing uses a four-wire Kelvin "
  "connection across a 10 mOhm shunt; charge and discharge are two independent "
  "paths, of which only one is energized at a time. The architecture is shown in "
  "Fig. 1.")

add_figure(os.path.join(FIG, "fig_arch_en.png"), 8.4,
           "Test-platform system architecture.")

p(S_H2, "Current calibration")
p(S_BODY,
  "Because every method rests on current integration, current accuracy is the "
  "foundation of the platform. The raw INA226 reading was systematically about "
  "15.8% high (mainly shunt tolerance and ADC offset). A 14-point piecewise-linear "
  "look-up table (seven charge and seven discharge points from 0 to 2 A) was built "
  "using the instruments as reference. Validated against four independent currents "
  "not in the table, the full-range error is below 0.21 mA, i.e. 0.012% at 1.75 A, "
  "down from the +13-16% raw bias.")

p(S_H2, "Cross-round test protocol")
p(S_BODY,
  "One round comprises four “charge - rest - discharge - rest” groups; "
  "charging is always 0.5C, and the four discharge rates are 0.5C, 1.0C, 1.5C and "
  "2.0C (1C = 2 A). A dV/dI perturbation is injected every 60 s during discharge - "
  "a brief step down to 0.2C - to serve the dynamic-impedance method, and Coulomb "
  "counting covers the perturbation seconds. Records persist across runs. The cell "
  "is a custom lithium-ion unit (2000 mAh, V_cv = 4.2 V, V_cut = 2.5 V); its rate "
  "capability is very flat, with discharge capacity dropping only about 1% from "
  "0.5C to 2.0C and good round-to-round repeatability, providing a favorable and "
  "stable basis for the comparison.")

p(S_H2, "The three estimators")
p(S_BODY,
  "All three run on a common firmware skeleton at a 1 Hz update; only the core "
  "estimation module is swapped, keeping the measurement path, timebase and I/O "
  "identical for fairness.")
p(S_BODY,
  "**Coulomb counting** recursively integrates the calibrated current as in (2), "
  "and also serves as the per-cycle ground truth after normalization by the "
  "measured full-discharge capacity q_full.")
p(S_BODY,
  "**Extended Kalman Filter.** With state x = [SOC, V1] and the first-order RC "
  "model, the observation is nonlinear only through V_OC(SOC), read from a 21-point "
  "GITT pseudo-OCV table refined to a 1% grid so its slope (the Jacobian) is "
  "continuous. Because there are two states and one scalar observation, the gain "
  "needs only a scalar division - no matrix inversion - which is the key to "
  "real-time operation on the MCU. Model parameters, identified by least squares "
  "from GITT relaxation pulses, are R0 = 51.9 mOhm (cell-terminal domain), "
  "R1 = 21.3 mOhm and tau_1 = 177.5 s.")
p(S_BODY,
  "**Dynamic impedance.** The instantaneous impedance Z = ΔV/ΔI from each "
  "perturbation is fit against SOC by the quadratic of (3), whose inverse yields "
  "SOC (with a branch selection because the parabola is symmetric). The combined "
  "fit is a = 20.2, b = -21.6, c = 63.6 mOhm with a minimum near 53% SOC, "
  "consistent with the literature [1].")

# -------------------------------------------------------------------- IV. Results
p(S_H1, "Experimental Results")
p(S_H2, "Accuracy")
p(S_BODY,
  "On one round of four discharge cycles, the three methods are scored against the "
  "per-cycle re-anchored Coulomb ground truth (Table I); Fig. 2 overlays the "
  "trajectories. The EKF tracks the truth almost exactly, with RMSE below 1% at "
  "every rate - the best of the three. Coulomb counting is a stable 1.7-2.0% with a "
  "small constant negative bias that traces to a scale difference between the board "
  "and bench current chains, not to the algorithm. Impedance-only online inversion "
  "is ill-posed for this flat-curve cell (23-31%): the whole Z-SOC curve spans only "
  "about 6 mOhm while single-event measurement noise is about 2 mOhm, so mid-SOC "
  "branch selection is nearly random. Under the idealized assumption of always "
  "choosing the correct branch, the offline RMSE is 6-10%, an upper bound not "
  "available online.")

# The spanning figure is emitted BEFORE Table I so that it lands at the top of the
# page across the full width. Placing the table first left the 2-column section
# holding nothing but the table, which Word balanced into the left column only and
# left the top-right quarter of the page blank.
add_figure(os.path.join(FIG, "fig_soc_en.png"), 17.2,
           "On-board SOC trajectories of the three methods versus the bench Coulomb "
           "ground truth (Round 41, four discharge rates; each panel annotates RMSE "
           "against the bench truth).", span=True)

add_table("Accuracy of the three methods (Round 41 board measurement; "
          "dynamic-impedance offline over rounds 1-3)",
          [["Method", "RMSE (%)", "MAE (%)", "e_max (%)", "Note"],
           ["Coulomb (board)", "1.7-2.0", "1.5-1.7", "3.0-3.5", "measurement-chain bias"],
           ["EKF (board)", "0.22-0.90", "0.20-0.76", "0.6-1.9", "best; <1% all rates"],
           ["Dyn. imp. (online)", "23-31", "-", "~96", "inversion ill-posed"],
           ["Dyn. imp. (offline)", "6-10", "5-8", "~20", "oracle-branch bound"]],
          [1230, 800, 760, 800, 1450])

p(S_H2, "Measurement-domain consistency (key finding)")
p(S_BODY,
  "Before correction (Round 40) the EKF showed a rate-dependent positive bias "
  "(RMSE 0.87-4.48%). The cause was a measurement-domain mismatch: R0 had been "
  "identified from the load-terminal voltage, which includes about 24 mOhm of "
  "wiring and contact resistance, whereas the board EKF uses the INA226 voltage at "
  "the cell terminals. Re-referencing R0 to the cell domain (51.9 mOhm) drove all "
  "four rates below 1% and removed the rate-dependent bias entirely (Round 41). The "
  "dynamic-impedance method exposes the same root cause more starkly: the "
  "bench-domain parabola minimum (57.9 mOhm) lies above every board measurement, so "
  "the discriminant has no real root and the whole table fails. Between the two "
  "domains only the constant term c shifts by about 24 mOhm - the quadratic and "
  "linear terms, set by the cell electrochemistry, are essentially unchanged [8]. "
  "The engineering lesson is that parameters and look-up tables must be built in "
  "the same measurement domain as deployment; otherwise a correct algorithm still "
  "fails systematically. This is a concrete instance of the "
  "evaluation-versus-deployment distortion raised in Section I.")

p(S_H2, "Robustness")
p(S_BODY,
  "Three stress tests (Table II) compare the methods. On a wrong initial value, "
  "Coulomb counting never recovers (mitigated here by an automatic full-charge "
  "re-anchor), the EKF is pulled back within a single measurement update, and "
  "dynamic impedance re-estimates at the first event. Under C-rate stepping the EKF "
  "absorbs the transient in its RC branch, whereas dynamic impedance mistakes the "
  "step for a perturbation event. Under injected sensor noise, Coulomb counting and "
  "the EKF stay bounded while impedance inversion flips branches catastrophically - "
  "confirming the ill-posedness is mathematical, not a sensor-quality issue.")

add_table("Robustness comparison (initial-value and C-rate rows are board "
          "measurements; noise row is a PC replay with injected Gaussian noise)",
          [["Stress test", "Coulomb", "EKF", "Dynamic impedance"],
           ["Wrong initial value", "no recovery", "single-step pull-back", "re-estimate <=60 s"],
           ["C-rate switch peak", "<=0.21%", "<=0.48%", "10-48%"],
           ["SOC jitter under noise", "<=0.054%", "<=0.27% (bounded)", "~53% (branch flip)"]],
          [1500, 1000, 1200, 1340])

p(S_H2, "Embedded footprint")
p(S_BODY,
  "Table III reports the extra Flash, RAM and per-update cycles each method adds "
  "over the shared skeleton on an STM32G071 at 64 MHz (-Og, software floating "
  "point). The EKF costs about 50x the compute of Coulomb counting, yet at 237 us "
  "its update is only 0.024% of the 1-second budget. Thus “EKF is too "
  "heavy” does not hold at this MCU tier; the resource argument only becomes "
  "decisive on lower-end platforms or at much higher update rates.")

add_table("Embedded resource cost, measured as the difference against the shared "
          "firmware skeleton (medians over ~159,300 updates)",
          [["Method", "Flash (B)", "RAM (B)", "Cycles/update", "FPU"],
           ["Coulomb", "+1,892", "+24", "305 (4.8 us)", "no"],
           ["EKF", "+2,336", "+40", "15,187 (237 us)", "soft"],
           ["Dyn. imp. (no event)", "+1,540", "+48", "442 (6.9 us)", "soft"],
           ["Dyn. imp. (event)", "same", "same", "6,525 (102 us)", "soft"]],
          [1450, 830, 730, 1350, 680])

# ---------------------------------------------------------------- V. Conclusions
p(S_H1, "Conclusions")
p(S_BODY,
  "On a single embedded platform with identical cell, protocol and MCU, this work "
  "implemented and fairly compared three SOC estimation methods. Three findings "
  "stand out. First, the EKF is the most accurate (0.22-0.90%, below 1% at all "
  "rates), Coulomb counting is a stable middle (1.7-2.0%), and impedance-only "
  "online inversion is infeasible (23-31%) for a flat Z-SOC cell; the accuracy "
  "ordering spans an order of magnitude and every error source is traced. Second, "
  "the resource ladder is quantified: the EKF costs about 50x Coulomb counting but "
  "still only 237 us per update, so it is not “too heavy” at this tier. "
  "Third, and most important, measurement-domain consistency dominates algorithm "
  "choice - parameters and tables built in a domain different from deployment fail "
  "systematically, as proven by the independent re-measurement that took the EKF "
  "from 4.48% to below 1%. Dynamic impedance remaining infeasible even after domain "
  "correction further shows that feasibility is set jointly by cell characteristics "
  "and the measurement chain, not by the algorithm alone. The reproducible "
  "fair-comparison platform is the main contribution; future work extends it to SOH "
  "estimation, temperature compensation [10] and LFP hysteresis [11].")

# ------------------------------------------------------------------ References
p(S_HEADX, "References", size=12, bold=True)
REFS = [
    "C.-H. Lin, C.-M. Wang, and C.-Y. Ho, “Implementation of state-of-charge "
    "and state-of-health estimation for lithium-ion batteries,” in Proc. IECON "
    "2016, Florence, Italy, Oct. 2016, pp. 18-24.",
    "A. Hasan, M. Skriver, and T. A. Johansen, “eXogenous Kalman filter for "
    "lithium-ion batteries state-of-charge estimation in electric vehicles,” "
    "in Proc. IEEE Conf. Control Technol. Appl. (CCTA), Copenhagen, Denmark, Aug. "
    "2018, pp. 1403-1408.",
    "A. Barros, P.-J. Kirsch, and J. Sun, “Adaptive extended Kalman filtering "
    "for battery state of charge estimation on STM32,” IEEE Embedded Syst. "
    "Lett., vol. 17, no. 3, pp. 121-124, 2025.",
    "K. Movassagh, A. Raihan, B. Balasingam, and K. Pattipati, “A critical "
    "look at Coulomb counting approach for state of charge estimation in "
    "batteries,” Energies, vol. 14, no. 14, art. 4074, 2021.",
    "I. Baccouche, S. Jemmali, B. Manai, N. Omar, and N. E. B. Amara, "
    "“Implementation of an improved Coulomb-counting algorithm based on a "
    "piecewise SOC-OCV relationship for SOC estimation of Li-ion battery,” "
    "Int. J. Renew. Energy Res., vol. 8, no. 1, 2018.",
    "S. Zhao and D. A. Howey, “Global sensitivity analysis of battery "
    "equivalent circuit model parameters,” in Proc. IEEE Veh. Power "
    "Propulsion Conf. (VPPC), Hangzhou, China, Oct. 2016, pp. 1-4.",
    "L. D. Couto and M. Kinnaert, “Partition-based unscented Kalman filter for "
    "reconfigurable battery pack state estimation using an electrochemical "
    "model,” in Proc. Amer. Control Conf. (ACC), Milwaukee, WI, USA, Jun. "
    "2018, pp. 3122-3127.",
    "A. Kulkarni, A. Barai, and J. Marco, “Novel low-complexity model "
    "development for Li-ion cells using online impedance measurement,” J. "
    "Energy Storage, vol. 91, art. 112034, 2024.",
    "J. Knox, M. Blyth, and A. Hales, “Advancing state estimation for "
    "lithium-ion batteries with hysteresis: systematic extended Kalman filter "
    "tuning,” Sci. Rep., vol. 14, art. 12472, 2024.",
    "Y. Qin, S. Adams, and C. Yuen, “A transfer learning-based state of charge "
    "estimation for lithium-ion battery at varying ambient temperatures,” IEEE "
    "Trans. Ind. Informat., vol. 17, no. 11, pp. 7304-7315, Nov. 2021.",
    "B. Yi, X. Du, J. Zhang, X. Wu, and Q. Hu, “Bias-compensated state of "
    "charge and state of health joint estimation for lithium iron phosphate "
    "batteries,” IEEE Trans. Transport. Electrific., early access, 2024.",
]
for r in REFS:
    p(S_REF, r)


# ------------------------------------------------------------------- equations
# The template's `equation` style carries tab stops (centre 2520, right 5040) and a
# Symbol font; force Times New Roman italic and use \t<eq>\t(n).
def eq_after(marker_text, text, number):
    """Insert a numbered display equation right after the paragraph containing marker."""
    target = None
    for par in doc.paragraphs:
        if par.style.name == S_BODY and marker_text in par.text:
            target = par
            break
    assert target is not None, f"marker not found: {marker_text!r}"
    par = doc.add_paragraph()
    target._p.addnext(par._p)
    par.style = S_EQ
    for chunk, ital in ((("\t" + text), True), ("\t(%d)" % number, False)):
        r = par.add_run(chunk)
        r.font.name = TIMES
        r.font.italic = ital
        rpr = r._element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts')
            rpr.insert(0, rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rf.set(qn(a), TIMES)
    return par


eq_after("The terminal voltage is given by (1)",
         "V_t = V_OC(SOC) - I·R0 - V1", 1)
eq_after("recursively integrates the calibrated current as in (2)",
         "SOC(t) = SOC(t0) - (1 / C_rated) ∫ I(τ) dτ", 2)
eq_after("fit against SOC by the quadratic of (3)",
         "Z = ΔV / ΔI = a·SOC² + b·SOC + c", 3)


doc.save(OUT_DOCX)
print("saved", OUT_DOCX)
print("sections:", len(doc.sections))
for i, s in enumerate(doc.sections):
    c = s._sectPr.find(qn('w:cols'))
    print(f"  [{i}] start={s.start_type} cols={c.get(qn('w:num')) if c is not None else 1}")
