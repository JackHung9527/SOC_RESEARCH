# -*- coding: utf-8 -*-
"""
Build the TANET conference paper (6 pages) in the format of the lab reference paper
(CYCU, Dept. of Information and Computer Engineering): A4, single-column title block +
two-column body, bilingual abstract (Chinese 摘要 in KaiTi + English Abstract),
numbered bold section headings, Times New Roman body @10pt / 18pt leading.

Content condensed from the author's thesis (English body, per instruction to fully
follow the reference paper's style).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
OUT_DOCX = os.path.join(HERE, "A_Comparative_Study_and_Embedded_Implementation_of_SOC_Estimation_Methods_for_Lithium-ion_Batteries.docx")

TIMES = "Times New Roman"
KAI = "標楷體"          # 標楷體 / DFKai-SB
BLACK = RGBColor(0, 0, 0)

# ---------------------------------------------------------------- font helper
def _set_fonts(run, ascii_font, ea_font):
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii_font)
    rf.set(qn('w:hAnsi'), ascii_font)
    rf.set(qn('w:cs'), ascii_font)
    rf.set(qn('w:eastAsia'), ea_font)


def _runs_from(text):
    """Split **bold** spans -> list of (chunk, is_bold)."""
    out, bold = [], False
    for i, chunk in enumerate(text.split("**")):
        if chunk:
            out.append((chunk, i % 2 == 1))
    return out


def para(doc, text, *, base="times", size=10.0, bold=False, italic=False,
         align="just", line=18.0, before=0.0, after=0.0, first_indent=0.0,
         hanging=0.0, keep=False):
    ea = KAI if base == "kai" else KAI            # eastAsia always KaiTi
    asc = TIMES
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = {"just": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER}[align]
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_indent:
        pf.first_line_indent = Cm(first_indent)
    if hanging:
        pf.left_indent = Cm(hanging)
        pf.first_line_indent = Cm(-hanging)
    if keep:
        pf.keep_with_next = True
    for chunk, b in _runs_from(text):
        r = p.add_run(chunk)
        r.font.size = Pt(size)
        r.font.bold = bold or b
        r.font.italic = italic
        _set_fonts(r, asc, ea if base == "kai" else TIMES)
    # for english paragraphs still map eastAsia to KaiTi so stray CJK looks right
    if base != "kai":
        for r in p.runs:
            r._element.get_or_add_rPr().find(qn('w:rFonts')).set(qn('w:eastAsia'), KAI)
    return p


# ---------------------------------------------------------------- section geom
def _set_page(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(1.9)


def _set_cols(section, num, space_cm=0.6):
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(int(space_cm * 567)))   # cm -> twips
    cols.set(qn('w:equalWidth'), "1")


# ---------------------------------------------------------------- tables
def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexc)
    tcPr.append(sh)


def _cell(cell, text, *, bold=False, size=9.0, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(11.5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    for chunk, b in _runs_from(text):
        r = p.add_run(chunk)
        r.font.size = Pt(size)
        r.font.bold = bold or b
        _set_fonts(r, TIMES, KAI)


def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:cantSplit'))


def add_captioned_table(doc, caption, rows, *, widths=None, size=9.0):
    """Caption placed ABOVE the table (reference/convention); caption + all rows
    kept together in one column via keep_with_next + cantSplit."""
    para(doc, caption, base="times", size=9.0, bold=True, align="left",
         line=12, before=5, after=2, keep=True)
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for ri, row in enumerate(rows):
        _cant_split(t.rows[ri])
        for ci, val in enumerate(row):
            al = "center" if (ri == 0 or ci > 0) else "left"
            _cell(t.rows[ri].cells[ci], val, bold=(ri == 0), size=size, align=al)
            if ri == 0:
                _shade(t.rows[ri].cells[ci], "DCE6F4")
            # keep every row with the next one so the table never splits a column
            for p in t.rows[ri].cells[ci].paragraphs:
                p.paragraph_format.keep_with_next = (ri < len(rows) - 1)
    if widths:
        for ci, w in enumerate(widths):
            for ri in range(len(rows)):
                t.rows[ri].cells[ci].width = Cm(w)
    para(doc, "", base="times", size=4, line=6, after=4)
    return t


def add_figure(doc, path, width_cm, caption, *, span=False):
    if span:
        s = doc.add_section(WD_SECTION.CONTINUOUS)
        _set_page(s); _set_cols(s, 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    para(doc, caption, base="times", size=9.5, bold=True, align="center",
         line=13, before=1, after=6)
    if span:
        s2 = doc.add_section(WD_SECTION.CONTINUOUS)
        _set_page(s2); _set_cols(s2, 2)


# ================================================================ BUILD
doc = Document()
# normal style baseline
doc.styles['Normal'].font.name = TIMES
doc.styles['Normal'].font.size = Pt(10)

s0 = doc.sections[0]
_set_page(s0)
_set_cols(s0, 1)          # title block: single column

# ---- Title ----
para(doc, "A Comparative Study and Embedded Implementation of "
          "SOC Estimation Methods for Lithium-ion Batteries",
     base="times", size=16, bold=True, align="center", line=20,
     before=6, after=10)

# ---- Authors / affiliation ----
para(doc, "Ta-Chia Hung, Wei-Kai Cheng", base="times", size=12, align="center",
     line=16, after=0)
para(doc, "Department of Information and Computer Engineering",
     base="times", size=12, align="center", line=16, after=0)
para(doc, "Chung Yuan Christian University", base="times", size=12,
     align="center", line=16, after=0)
para(doc, "Taoyuan City, Taiwan", base="times", size=12, align="center",
     line=16, after=0)
para(doc, "g11277602@cycu.edu.tw, wkcheng@cycu.edu.tw", base="times", size=12,
     align="center", line=16, after=6)

# ---- switch to two columns ----
sb = doc.add_section(WD_SECTION.CONTINUOUS)
_set_page(sb)
_set_cols(sb, 2)

# ---- Chinese abstract ----
para(doc, "摘要", base="kai", size=12, bold=True, align="center",
     line=18, before=2, after=4)
CN_ABS = ("荷電狀態（State of Charge, SOC）估測是"
"電池管理系統的核心機能，然而既有"
"方法比較多在離線環境以單一精度"
"指標進行，鮮少在實際部署的資源"
"受限微控制器上量化精度與資源代"
"價，且常因電池、協定與硬體不一"
"致而缺乏公平基準。本研究建構一"
"套「同電池、同協定、同微控制器"
"」的自動化跨輪測試平台，並於同"
"一顆 STM32 上以一致韌體規範實作庫倫"
"計數、擴展卡爾曼濾波器（EKF）與"
"動態阻抗三種方法，於精度、強健"
"性與嵌入式資源三軸進行公平比較"
"。整輪板端實測顯示：EKF 於四種放"
"電倍率之均方根誤差均低於 1%"
"（0.22–0.90%）為三法最佳，庫倫計"
"數約 1.7–2.0%，動態阻抗線上單獨"
"反推因曲線平坦而病態（23–31%"
"）；資源代價 EKF 約為庫倫的 50 倍，"
"但於 64 MHz 平台僅佔每秒更新週期的萬"
"分之三。本研究並揭示一關鍵工程"
"發現：參數與對照表若非與部署端"
"量測鏈同域建立，演算法再正確亦"
"會系統性失準；量測域一致性之重"
"要性凌駕於演算法選擇本身。")
para(doc, CN_ABS, base="kai", size=10, align="just", line=18, first_indent=0.6)
para(doc, "關鍵詞：鋰離子電池、荷電狀"
          "態估測、庫倫計數、擴展卡爾"
          "曼濾波器、動態阻抗法、嵌入"
          "式系統", base="kai", size=10, align="just", line=18, after=6)

# ---- English abstract ----
para(doc, "Abstract", base="times", size=12, bold=True, align="center",
     line=18, before=2, after=4)
EN_ABS = ("State of Charge (SOC) estimation is a core function of a battery "
"management system, yet existing method comparisons are mostly performed offline "
"using a single accuracy metric and rarely quantify, on the resource-constrained "
"microcontroller where the algorithm is actually deployed, both accuracy and "
"resource cost; differing cells, protocols and hardware further prevent a fair "
"baseline. This work builds an automated multi-round test platform that keeps the "
"cell, the test protocol and the microcontroller identical across methods, and "
"implements Coulomb counting, an Extended Kalman Filter (EKF) and a dynamic-"
"impedance method on the same STM32 under a unified firmware specification. "
"Full-round on-board measurements show that the EKF attains a root-mean-square "
"error below 1% (0.22-0.90%) at all four discharge rates - the best of the three - "
"while Coulomb counting yields 1.7-2.0% and impedance-only online inversion is "
"ill-posed (23-31%) for this flat Z-SOC cell. The EKF costs about 50x the compute "
"of Coulomb counting yet consumes only 0.024% of the 1-second update budget at "
"64 MHz. We further show that measurement-domain consistency dominates algorithm "
"choice: parameters and look-up tables built in a different measurement domain "
"than deployment cause systematic error regardless of algorithm correctness.")
para(doc, EN_ABS, base="times", size=10, align="just", line=18, first_indent=0.4)
para(doc, "**Keywords**: lithium-ion battery, state of charge estimation, Coulomb "
          "counting, extended Kalman filter, dynamic impedance, embedded systems",
     base="times", size=10, align="just", line=18, after=6)


# ---- helpers for headings ----
def h1(n, title):
    para(doc, f"{n}.  {title}", base="times", size=12, bold=True, align="left",
         line=18, before=9, after=4, keep=True)


def h2(n, title):
    para(doc, f"{n}  {title}", base="times", size=11, bold=True, align="left",
         line=18, before=6, after=3, keep=True)


def body(text, indent=0.4):
    para(doc, text, base="times", size=10, align="just", line=18,
         first_indent=indent, after=0)


def eq(text):
    para(doc, text, base="times", size=10.5, italic=True, align="center",
         line=18, before=3, after=3)


# ============================================================ 1. Introduction
h1(1, "Introduction")
body("Driven by decarbonization and the electrification of transport, lithium-ion "
"batteries have become the dominant energy-storage element in electric vehicles, "
"portable electronics and grid storage. The battery management system (BMS) is "
"therefore indispensable, and among its functions the State of Charge (SOC) - the "
"ratio of remaining to rated capacity - plays the role of a fuel gauge for range "
"prediction, energy scheduling and over-charge/over-discharge protection. SOC is "
"an internal state that cannot be sensed directly; it can only be inferred from "
"terminal voltage, current and temperature of a highly nonlinear, time-varying "
"electrochemical cell [1], [4].", indent=0.0)

body("Although comparisons of SOC estimation methods are abundant, two gaps remain "
"insufficiently addressed. The first is evaluation-environment distortion. Most "
"comparisons are conducted on a PC or in MATLAB and report the root-mean-square "
"error (RMSE) as almost the only metric, whereas the real deployment target is a "
"resource-constrained microcontroller unit (MCU). How much Flash and RAM a method "
"occupies after porting, how many CPU cycles each update costs, and how much "
"accuracy is lost after fixed-point implementation are seldom quantified. Notably, "
"commercial BMS ICs generally adopt Coulomb counting with OCV correction rather "
"than the theoretically more accurate Extended Kalman Filter (EKF) [5] - an "
"engineering trade-off among memory, computation and cost that is rarely presented "
"with measured data.")

body("The second gap is unfair benchmarking. Existing cross-method comparisons "
"often use different cells, protocols and hardware, so the reported accuracies "
"share no common baseline. The dynamic-impedance method [1], for instance, claims "
"to need no initial value and to suit embedded use, yet it has never been placed on "
"the same MCU, the same cell and the same protocol as an EKF for an equal "
"resource-versus-accuracy comparison.")

body("To address both gaps, this work builds a \"same cell, same protocol, same "
"MCU\" fair-comparison platform and quantifies the accuracy-cost trade-off of "
"embedded SOC estimation. Its contributions are: (1) an automated, long-term "
"unattended cross-round test platform integrating an STM32, a programmable DC "
"supply, a programmable electronic load and a high-accuracy current/voltage sensor "
"with an automation scheduler; (2) an implementation of Coulomb counting, the EKF "
"and the dynamic-impedance method on one STM32 under a unified firmware "
"specification; and (3) an equal, reproducible comparison of accuracy, robustness "
"and embedded resource cost, from which we derive both a method-selection guide and "
"a key finding that measurement-domain consistency dominates algorithm choice.")

# ============================================================ 2. Background
h1(2, "Background and Related Work")
h2("2.1", "Equivalent-circuit model")
body("A first-order RC (Thevenin) equivalent-circuit model is adopted as the cell "
"model: an open-circuit-voltage source V_OC(SOC) in series with an ohmic resistance "
"R0 and one parallel RC branch (R1, C1) that captures polarization. The terminal "
"voltage is V_t = V_OC(SOC) - I*R0 - V1, where V1 is the RC-branch voltage. This "
"structure balances fidelity and identifiability for embedded use [6], and the "
"nonlinear V_OC(SOC) relation is obtained by a pseudo-OCV table [5].", indent=0.0)

h2("2.2", "SOC estimation methods")
body("Coulomb counting integrates current to accumulate charge; it is simple and "
"model-free but has no self-correction and accumulates drift [4]. OCV look-up maps "
"a rested open-circuit voltage to SOC but requires long relaxation. Model-based "
"methods, dominated by the Kalman-filter family (KF/EKF/UKF), fuse current "
"integration with a voltage measurement to correct the estimate and can recover "
"from a wrong initial value [2], [3], [7], [9]. The dynamic-impedance method [1] "
"estimates SOC from the instantaneous voltage-to-current ratio during operation, "
"avoiding both an initial value and long rest, and online impedance has also been "
"used for low-complexity modeling and aging indication [8]. Commercial ICs "
"nonetheless favor Coulomb counting with OCV correction for resource reasons [5], a "
"trade-off this work quantifies on real hardware.")

# ============================================================ 3. Platform
h1(3, "Test Platform and Estimator Implementation")
h2("3.1", "System architecture")
body("The platform plays two roles simultaneously: an excitation-and-ground-truth "
"rig and an embedded estimation target. A programmable DC supply (ITECH IT6302, "
"CC-CV) and electronic load (ITECH IT8512A+, CC) apply controlled charge/discharge "
"to a single device-under-test (DUT), while an STM32G071 with an INA226 "
"current/voltage sensor runs the estimators. Both roles share the same cell and the "
"same high-power loop, so the instrument ground truth and the MCU estimate can be "
"aligned point-by-point. Sensing uses a four-wire Kelvin connection across a 10 mOhm "
"shunt; charge and discharge are two independent paths, of which only one is "
"energized at a time. The architecture is shown in Fig. 1.", indent=0.0)

add_figure(doc, os.path.join(FIG, "fig_arch_en.png"), 7.7,
           "Figure 1.  Test-platform system architecture.")

h2("3.2", "Current calibration")
body("Because every method rests on current integration, current accuracy is the "
"foundation of the platform. The raw INA226 reading was systematically about 15.8% "
"high (mainly shunt tolerance and ADC offset). A 14-point piecewise-linear "
"look-up table (seven charge and seven discharge points from 0 to 2 A) was built "
"using the instruments as reference. Validated against four independent currents "
"not in the table, the full-range error is below 0.21 mA, i.e. 0.012% at 1.75 A, "
"down from the +13-16% raw bias.", indent=0.0)

h2("3.3", "Cross-round test protocol")
body("One round comprises four \"charge - rest - discharge - rest\" groups; "
"charging is always 0.5C, and the four discharge rates are 0.5C, 1.0C, 1.5C and "
"2.0C (1C = 2 A). A dV/dI perturbation is injected every 60 s during discharge - a "
"brief step down to 0.2C - to serve the dynamic-impedance method, and Coulomb "
"counting covers the perturbation seconds. Records persist across runs. The cell is "
"a custom NMC unit (2000 mAh, V_cv = 4.2 V, V_cut = 2.5 V); its rate capability is "
"very flat, with discharge capacity dropping only about 1% from 0.5C to 2.0C and "
"good round-to-round repeatability, providing a favorable and stable basis for the "
"comparison.", indent=0.0)

h2("3.4", "The three estimators")
body("All three run on a common firmware skeleton at a 1 Hz update; only the core "
"estimation module is swapped, keeping the measurement path, timebase and I/O "
"identical for fairness.", indent=0.0)
body("**Coulomb counting** recursively integrates the calibrated current, and also "
"serves as the per-cycle ground truth after normalization by the measured "
"full-discharge capacity q_full:")
eq("SOC(t) = SOC(t0) - (1 / C_rated) ∫ I(τ) dτ")
body("**Extended Kalman Filter.** With state x = [SOC, V1] and the first-order RC "
"model, the observation is nonlinear only through V_OC(SOC), read from a 21-point "
"GITT pseudo-OCV table refined to a 1% grid so its slope (the Jacobian) is "
"continuous. Because there are two states and one scalar observation, the gain "
"needs only a scalar division - no matrix inversion - which is the key to real-time "
"operation on the MCU. Model parameters, identified by least squares from GITT "
"relaxation pulses, are R0 = 51.9 mOhm (cell-terminal domain), R1 = 21.3 mOhm and "
"τ1 = 177.5 s.")
eq("V_t = V_OC(SOC) - I·R0 - V1")
body("**Dynamic impedance.** The instantaneous impedance Z = ΔV/ΔI from "
"each perturbation is fit against SOC by a quadratic, whose inverse yields SOC "
"(with a branch selection because the parabola is symmetric). The combined fit is "
"a = 20.2, b = -21.6, c = 63.6 mOhm with a minimum near 53% SOC, consistent with "
"the literature [1].")
eq("Z = ΔV / ΔI = a·SOC² + b·SOC + c")

# ============================================================ 4. Results
h1(4, "Experimental Results")
h2("4.1", "Accuracy")
body("On one round of four discharge cycles, the three methods are scored against "
"the per-cycle re-anchored Coulomb ground truth (Table 1); Fig. 2 overlays the "
"trajectories. The EKF tracks the truth almost exactly, with RMSE below 1% at every "
"rate - the best of the three. Coulomb counting is a stable 1.7-2.0% with a small "
"constant negative bias that traces to a scale difference between the board and "
"bench current chains, not to the algorithm. Impedance-only online inversion is "
"ill-posed for this flat-curve cell (23-31%): the whole Z-SOC curve spans only "
"about 6 mOhm while single-event measurement noise is about 2 mOhm, so mid-SOC "
"branch selection is nearly random. Under the idealized assumption of always "
"choosing the correct branch, the offline RMSE is 6-10%, an upper bound not "
"available online.", indent=0.0)

add_captioned_table(doc,
    "Table 1.  Accuracy of the three methods (Round 41 board measurement; "
    "dynamic-impedance offline over rounds 1-3).",
    [
    ["Method", "RMSE (%)", "MAE (%)", "e_max (%)", "Note"],
    ["Coulomb (board)", "1.7-2.0", "1.5-1.7", "3.0-3.5", "measurement-chain bias"],
    ["EKF (board)", "0.22-0.90", "0.20-0.76", "0.6-1.9", "best; <1% all rates"],
    ["Dyn. imp. (online)", "23-31", "-", "~96", "inversion ill-posed"],
    ["Dyn. imp. (offline)", "6-10", "5-8", "~20", "oracle-branch bound"],
    ], widths=[2.5, 1.5, 1.4, 1.4, 2.4], size=8.6)

add_figure(doc, os.path.join(FIG, "fig_soc_en.png"), 15.6,
           "Figure 2.  On-board SOC trajectories of the three methods versus the "
           "bench Coulomb ground truth (Round 41, four discharge rates; each panel "
           "annotates RMSE against the bench truth).", span=True)

h2("4.2", "Measurement-domain consistency (key finding)")
body("Before correction (Round 40) the EKF showed a rate-dependent positive bias "
"(RMSE 0.87-4.48%). The cause was a measurement-domain mismatch: R0 had been "
"identified from the load-terminal voltage, which includes about 24 mOhm of wiring "
"and contact resistance, whereas the board EKF uses the INA226 voltage at the cell "
"terminals. Re-referencing R0 to the cell domain (51.9 mOhm) drove all four rates "
"below 1% and removed the rate-dependent bias entirely (Round 41). The dynamic-"
"impedance method exposes the same root cause more starkly: the bench-domain "
"parabola minimum (57.9 mOhm) lies above every board measurement, so the "
"discriminant has no real root and the whole table fails. Between the two domains "
"only the constant term c shifts by about 24 mOhm - the quadratic and linear terms, "
"set by the cell electrochemistry, are essentially unchanged [8]. The engineering "
"lesson is that parameters and look-up tables must be built in the same measurement "
"domain as deployment; otherwise a correct algorithm still fails systematically. "
"This is a concrete instance of the evaluation-versus-deployment distortion raised "
"in Section 1.", indent=0.0)

h2("4.3", "Robustness")
body("Three stress tests (Table 2) compare the methods. On a wrong initial value, "
"Coulomb counting never recovers (mitigated here by an automatic full-charge "
"re-anchor), the EKF is pulled back within a single measurement update, and dynamic "
"impedance re-estimates at the first event. Under C-rate stepping the EKF absorbs "
"the transient in its RC branch, whereas dynamic impedance mistakes the step for a "
"perturbation event. Under injected sensor noise, Coulomb counting and the EKF stay "
"bounded while impedance inversion flips branches catastrophically - confirming the "
"ill-posedness is mathematical, not a sensor-quality issue.", indent=0.0)

add_captioned_table(doc,
    "Table 2.  Robustness comparison (initial-value and C-rate rows are board "
    "measurements; noise row is a PC replay with injected Gaussian noise).",
    [
    ["Stress test", "Coulomb", "EKF", "Dynamic impedance"],
    ["Wrong initial value", "no recovery", "single-step pull-back", "re-estimate <=60 s"],
    ["C-rate switch peak", "<=0.21%", "<=0.48%", "10-48%"],
    ["SOC jitter under noise", "<=0.054%", "<=0.27% (bounded)", "~53% (branch flip)"],
    ], widths=[3.1, 1.9, 2.3, 2.4], size=8.6)

h2("4.4", "Embedded footprint")
body("Table 3 reports the extra Flash, RAM and per-update cycles each method adds "
"over the shared skeleton on an STM32G071 at 64 MHz (-Og, software floating point). "
"The EKF costs about 50x the compute of Coulomb counting, yet at 237 us its update "
"is only 0.024% of the 1-second budget. Thus \"EKF is too heavy\" does not hold at "
"this MCU tier; the resource argument only becomes decisive on lower-end platforms "
"or at much higher update rates.", indent=0.0)

add_captioned_table(doc,
    "Table 3.  Embedded resource cost, measured as the difference against the "
    "shared firmware skeleton (medians over ~159,300 updates).",
    [
    ["Method", "Flash (B)", "RAM (B)", "Cycles/update", "FPU"],
    ["Coulomb", "+1,892", "+24", "305 (4.8 us)", "no"],
    ["EKF", "+2,336", "+40", "15,187 (237 us)", "soft"],
    ["Dyn. imp. (no event)", "+1,540", "+48", "442 (6.9 us)", "soft"],
    ["Dyn. imp. (event)", "same", "same", "6,525 (102 us)", "soft"],
    ], widths=[3.0, 1.6, 1.4, 2.3, 1.4], size=8.6)

# ============================================================ 5. Conclusions
h1(5, "Conclusions")
body("On a single embedded platform with identical cell, protocol and MCU, this "
"work implemented and fairly compared three SOC estimation methods. Three findings "
"stand out. First, the EKF is the most accurate (0.22-0.90%, below 1% at all "
"rates), Coulomb counting is a stable middle (1.7-2.0%), and impedance-only online "
"inversion is infeasible (23-31%) for a flat Z-SOC cell; the accuracy ordering "
"spans an order of magnitude and every error source is traced. Second, the resource "
"ladder is quantified: the EKF costs about 50x Coulomb counting but still only "
"237 us per update, so it is not \"too heavy\" at this tier. Third, and most "
"important, measurement-domain consistency dominates algorithm choice - parameters "
"and tables built in a domain different from deployment fail systematically, as "
"proven by the independent re-measurement that took the EKF from 4.48% to below 1%. "
"Dynamic impedance remaining infeasible even after domain correction further shows "
"that feasibility is set jointly by cell characteristics and the measurement chain, "
"not by the algorithm alone. The reproducible fair-comparison platform is the main "
"contribution; future work extends it to SOH estimation, temperature compensation "
"[10] and LFP hysteresis [11].", indent=0.0)

# ============================================================ 6. References
h1(6, "References")
REFS = [
"C.-H. Lin, C.-M. Wang, and C.-Y. Ho, \"Implementation of state-of-charge and "
"state-of-health estimation for lithium-ion batteries,\" in Proc. IECON 2016, "
"Florence, Italy, Oct. 2016, pp. 18-24.",
"A. Hasan, M. Skriver, and T. A. Johansen, \"eXogenous Kalman filter for "
"lithium-ion batteries state-of-charge estimation in electric vehicles,\" arXiv:"
"1810.09014, 2018.",
"A. Barros et al., \"Adaptive extended Kalman filtering for battery state of "
"charge estimation on STM32,\" arXiv:2504.05936, 2025.",
"K. Movassagh, S. A. Raihan, B. Balasingam, and K. Pattipati, \"A critical look "
"at Coulomb counting towards improving the Kalman-filter-based state-of-charge "
"tracking algorithms in rechargeable batteries,\" arXiv:2101.05435, 2021.",
"I. Baccouche et al., \"Implementation of an improved Coulomb-counting algorithm "
"based on a piecewise SOC-OCV relationship for SOC estimation of Li-ion battery,\" "
"arXiv:1803.10654, 2018.",
"S. Zhao and D. A. Howey, \"Global sensitivity analysis of battery equivalent "
"circuit model parameters,\" arXiv:1604.01293, 2016.",
"L. D. Couto and M. Kinnaert, \"Partition-based unscented Kalman filter for "
"reconfigurable battery pack state estimation using an electrochemical model,\" "
"arXiv:1709.07816, 2017.",
"A. Kulkarni et al., \"Novel low-complexity model development for Li-ion cells "
"using online impedance measurement,\" arXiv:2402.07777, 2024.",
"J. Knox, M. Blyth, and A. Hales, \"Advancing state estimation for lithium-ion "
"batteries with hysteresis: systematic extended Kalman filter tuning,\" arXiv:"
"2311.16942, 2023.",
"Y. Qin, S. Adams, and C. Yuen, \"A transfer-learning-based state of charge "
"estimation for lithium-ion battery at varying ambient temperatures,\" IEEE Trans. "
"Ind. Informat., 2021 (arXiv:2101.03704).",
"B. Yi et al., \"Bias-compensated state of charge and state of health joint "
"estimation for lithium iron phosphate batteries,\" arXiv:2401.08136, 2024.",
]
for i, r in enumerate(REFS, 1):
    para(doc, f"[{i}]  {r}", base="times", size=9.5, align="just", line=13,
         hanging=0.7, after=2)

doc.save(OUT_DOCX)
print("saved", OUT_DOCX)
print("sections:", len(doc.sections),
      "cols:", [s._sectPr.find(qn('w:cols')).get(qn('w:num')) if s._sectPr.find(qn('w:cols')) is not None else '?' for s in doc.sections])
