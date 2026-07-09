#!/usr/bin/env python3
"""sync_ch4_docx_20260707.py — 將第四章 Round 40 實測進版同步進論文 docx。

_20260702.docx → _20260707.docx，僅動第四章（元素 322–430 範圍內），
全部為文字段落與表格，不觸碰圖片/圖說/其他章節。

段落改寫：deepcopy 原段元素保 pPr，run 以原首 run 之 rPr 為樣板；
支援 **粗體** 與 X_y 下標（渲染成真 subscript run，比照 2026-06-25 之
底線變數轉換慣例）。表格：同形狀者逐格改字、需加列者 deepcopy 末列。
"""

import copy
import re
import shutil
import sys

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

SRC = "DOC/論文撰寫/鋰電池SOC估測方法之比較與嵌入式實作_20260702.docx"
DST = "DOC/論文撰寫/鋰電池SOC估測方法之比較與嵌入式實作_20260707.docx"

# ---------- 渲染：**bold** 與 base_sub 下標 ----------
TOK = re.compile(r"\*\*(.+?)\*\*|([A-Za-zΔτ]+)_([A-Za-z0-9max]+)")


def _segments(text):
    """→ [(text, bold, sub)]"""
    out, pos = [], 0
    for m in TOK.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], False, False))
        if m.group(1) is not None:
            # 粗體內仍可能含下標
            inner = m.group(1)
            p2 = 0
            for m2 in re.finditer(r"([A-Za-zΔτ]+)_([A-Za-z0-9max]+)", inner):
                if m2.start() > p2:
                    out.append((inner[p2:m2.start()], True, False))
                out.append((m2.group(1), True, False))
                out.append((m2.group(2), True, True))
                p2 = m2.end()
            if p2 < len(inner):
                out.append((inner[p2:], True, False))
        else:
            out.append((m.group(2), False, False))
            out.append((m.group(3), False, True))
        pos = m.end()
    if pos < len(text):
        out.append((text[pos:], False, False))
    return [s for s in out if s[0]]


def _mk_run(template_r, text, bold, sub):
    r = copy.deepcopy(template_r)
    # 清 text 節點與既有 br/tab
    for child in list(r):
        if child.tag != qn("w:rPr"):
            r.remove(child)
    t = r.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
    t.text = text
    r.append(t)
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    for b in rPr.findall(qn("w:b")):
        rPr.remove(b)
    if bold:
        rPr.append(rPr.makeelement(qn("w:b"), {}))
    for va in rPr.findall(qn("w:vertAlign")):
        rPr.remove(va)
    if sub:
        va = rPr.makeelement(qn("w:vertAlign"), {})
        va.set(qn("w:val"), "subscript")
        rPr.append(va)
    return r


def set_p(p, text):
    """整段換字，保 pPr 與字型（以原首 run 為樣板）。"""
    runs = p._p.findall(qn("w:r"))
    if not runs:
        sys.exit(f"段落無 run 可當樣板: {p.text[:30]}")
    template = runs[0]
    segs = _segments(text)
    anchor = template
    new_runs = [_mk_run(template, s, b, v) for s, b, v in segs]
    for r in runs:
        if r is not template:
            p._p.remove(r)
    for nr in new_runs:
        anchor.addnext(nr)
        anchor = nr
    p._p.remove(template)


def insert_after(p_anchor, text):
    """以 p_anchor 為樣板，複製一段插於其後並填字。回傳新 Paragraph。"""
    new_el = copy.deepcopy(p_anchor._p)
    p_anchor._p.addnext(new_el)
    np = Paragraph(new_el, p_anchor._parent)
    set_p(np, text)
    return np


def set_cell(cell, text):
    p = cell.paragraphs[0]
    runs = p._p.findall(qn("w:r"))
    if not runs:  # 空格：借同列他格 run 樣板
        for c2 in cell._parent.cells if hasattr(cell, "_parent") else []:
            pass
        # 直接建陽春 run
        r = p._p.makeelement(qn("w:r"), {})
        p._p.append(r)
        t = r.makeelement(qn("w:t"), {})
        t.text = text
        r.append(t)
        return
    set_p(p, text)
    for extra in cell.paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def main():
    shutil.copyfile(SRC, DST)
    d = docx.Document(DST)

    # ---- 第四章範圍內以前綴找段落 ----
    ch4 = []
    in4 = False
    for el in d.element.body:
        if el.tag == qn("w:p"):
            p = Paragraph(el, d)
            t = p.text.strip()
            if t.startswith("第四章") and p.style.name == "Heading 1":
                in4 = True
            elif t.startswith("第五章") and p.style.name == "Heading 1":
                in4 = False
            if in4:
                ch4.append(p)

    def P(prefix):
        for p in ch4:
            if p.text.strip().startswith(prefix):
                return p
        sys.exit(f"找不到段落: {prefix}")

    # ---- 表格（依表頭辨識，只在文件 tables 中找） ----
    def T(*head_prefixes):
        for tb in d.tables:
            heads = [c.text.strip() for c in tb.rows[0].cells]
            if all(any(h.startswith(x) for h in heads) for x in head_prefixes):
                return tb
        sys.exit(f"找不到表格: {head_prefixes}")

    # ============ 4.2 ============
    set_p(P("本節原理、狀態空間、演算法與嵌入式實作為設計完成"),
          "本節 (a) GITT 開路電壓對照表、(b) 模型參數辨識與 PC 端重放驗證、"
          "(c) 嵌入式移植與整輪板端實測，均已完成（4.2.5）。")

    set_p(P("此問題的狀態只有兩維、觀測只有一維"),
          "此問題的狀態只有兩維、觀測只有一維，因此增益計算僅需一次純量除法、無需一般矩陣求逆"
          "——這是擴展卡爾曼濾波器能在微控制器上即時運行的關鍵。其中電池內阻與極化等模型參數由 "
          "GITT 各脈衝之鬆弛曲線以最小平方辨識（SOC 20–90% 中段平均，n=15）：歐姆內阻 R_0 = "
          "75.9 mΩ（台架量測域；部署域換算見 4.2.5）、極化電阻 R_1 = 21.3 mΩ、時間常數 "
          "τ_1 = 177.5 s，逐脈衝擬合殘差均低於 2 mV。過程雜訊與量測雜訊以 PC 重放做網格搜尋，"
          "Q_SOC=10⁻⁸、R_meas=(10 mV)² 已位於最佳區（重放 RMSE 對 12 組 Q/R 組合變動不足 "
          "0.7%，顯示濾波器對調參不敏感）。")

    set_p(P("4.2.5 實驗結果"), "4.2.5 實驗結果（實測）")

    set_p(P("以下結果依賴 GITT OCV 表"),
          "**OCV 表（前置產物）**：GITT 於 2026-07-05 執行（放電單向、5% 步進 × 0.5C × "
          "30 分鐘鬆弛），取得含滿電錨點共 20 個平衡電位點（SOC 4.88–100%）；末步觸及 2.5 V "
          "化學下限提前中止，故 SOC < 4.88% 無實測點，表格 0% 節點為邊界外插並於韌體中如實註記。"
          "各點鬆弛末段 |dV/dt| 除最低 SOC 點外均低於 0.11 mV/min，平衡性充分。經 1% 細化、"
          "5 點移動平均後重取 5% 節點得 21 點對照表，全表嚴格單調遞增（分段線性雅可比恆正）。")

    set_p(P("OCV 表（前置產物）：[待測]"),
          "**PC 重放驗證**：以與韌體逐行等價的 PC 原型重放 GITT 全程資料（8,625 樣本、約 "
          "11.5 小時），以電流積分之庫倫 SOC 為真值：追蹤 RMSE 0.60%、最大誤差 2.2%；以電壓"
          "播種（開機情境）結果相同；刻意給 50% 錯誤初值時，因初始不確定度設為 (20%)²，首次量測"
          "更新即將估值拉回電壓所指之 SOC（單步收斂）。此步驟同時確認 Q/R 之選擇（見 4.2.3）。"
          "需註明此重放與參數辨識使用同一份 GITT 資料（in-sample），獨立驗證由下述整輪板端實測承擔。")

    set_p(P("標稱工況精度：對 0.5C"),
          "**板端整輪實測（Round 40，四倍率各一放電 cycle）**：STM32 端 EKF 每秒以 INA226 "
          "量測更新，與台架庫倫真值逐秒比對：")

    tb2 = T("倍率", "RMSE", "收斂時間")
    hdr = ["倍率", "RMSE (%)", "MAE (%)", "e_max (%)", "偏差方向"]
    for j, h in enumerate(hdr):
        set_cell(tb2.rows[0].cells[j], h)
    data2 = [["0.5C", "0.87", "0.85", "1.06", "+0.83"],
             ["1.0C", "2.50", "2.36", "4.00", "+2.36"],
             ["1.5C", "3.67", "3.45", "5.22", "+3.45"],
             ["2.0C", "4.48", "4.20", "6.44", "+4.20"]]
    for i, row in enumerate(data2, start=1):
        for j, v in enumerate(row):
            set_cell(tb2.rows[i].cells[j], v)

    set_p(P("表 4-2　EKF 各倍率精度"),
          "表 4-2　EKF 各倍率板端實測精度（Round 40；開機以電壓播種、起始即收斂，"
          "故不另列收斂時間；初值錯誤收斂見 4.4.2）")

    p_seed = P("初值錯誤收斂：刻意以")
    set_p(p_seed,
          "誤差隨倍率增大且恆為正偏，其來源已定位為**量測域錯位**：R_0 由 GITT 辨識時所用之"
          "負載電壓取自電子負載端（含電池至負載間之線材與接點電阻約 24 mΩ），而板端 EKF 使用 "
          "INA226 於電池端子之電壓——模型過電位因此被多算 24 mΩ × I，隨電流線性放大並將估值"
          "推高，與表中偏差趨勢一致（此量測域議題之完整分析見 4.3.4，兩法互為佐證）。將 R_0 "
          "換算至電池端域（51.9 mΩ）後之複測列為 [待測]（Round 41）。")
    insert_after(p_seed,
                 "**初值錯誤收斂（板端實測）**：透過除錯介面將運行中之 EKF 強制設為錯誤 SOC"
                 "（偏移 +25%），估值於一次量測更新內即被電壓項拉回、數秒內回到原軌跡——與 PC "
                 "重放之單步收斂行為一致，證實 EKF 相對庫倫計數的核心優勢（後者無任何修正機制，"
                 "見 4.4.2）。")

    # ============ 4.3 ============
    set_p(P("本節原理、演算法與嵌入式實作為設計完成"),
          "本節 4.3.4 之 Δ V/Δ I–SOC 二次擬合係數與反推 SOC 逐點精度為**實測**"
          "（fresh-cell rounds 1–3）；韌體移植與整輪線上驗證亦已完成（Round 40），"
          "並揭示重要的量測域發現（4.3.4 末）。")

    p404 = P("此精度水準說明：單以動態阻抗反推")
    set_p(p404,
          "此精度水準說明：**單以動態阻抗反推 SOC 不足以獨立提供高解析度估測，尤其在 SOC 中段**。"
          "這正當化了 4.3.3 所述之混合策略——以動態阻抗在擾動事件處提供「無須初值」的離散校正，"
          "事件之間再由庫倫計數內插以維持連續且高解析度的 SOC。")
    pa = insert_after(p404,
                 "**嵌入式部署之量測域發現（Round 40 實測）**：將表 4-3 之台架擬合係數移植入"
                 "微控制器後，於整輪測試中線上估測完全失效（RMSE 約 28%）。逐事件解剖原始樣本後"
                 "確認：板端量測本身乾淨（擾動 dwell 段有 2–3 個穩定樣本），但板端量得之動態阻抗"
                 "中位數僅 35.0 mΩ，遠低於台架同批擾動之 60–61 mΩ。差值約 24 mΩ 為**兩個量測點"
                 "之間的線材與接點電阻**——電子負載於其端子看阻抗（含配線），INA226 則在電池端子"
                 "直接量測，後者才是電芯真實動態阻抗。由於台架域係數之拋物線最低點（57.9 mΩ）高於"
                 "板端一切量測值，反解之判別式恆無實根，導致整組查表失效。")
    pb = insert_after(pa,
                 "以板端 447 個擾動事件對台架真值重新擬合，得電池端域係數 a=18.0、b=−21.1、"
                 "c=39.6 mΩ（擬合殘差 2.5 mΩ、最低點 SOC 58.6%）。與台架域係數（a=20.2、"
                 "b=−21.6、c=63.6）比較，**二次項與一次項幾乎不變、僅常數項平移約 −24 mΩ**："
                 "拋物線形狀由電芯電化學決定，常數項則由量測鏈決定。此發現之工程意涵為：**阻抗-SOC "
                 "對照表必須與部署端量測鏈同域建立**（或顯式校正量測點偏移），「離線建表、線上查表」"
                 "跨量測鏈搬移會使整組對照失效——此為評估環境與部署環境失真的具體實例，直接支持"
                 "第一章 1.2 節之研究缺口論述。附帶地，本輪亦發現擾動 dwell 實際長度依賴儀器命令"
                 "延遲（設定 1 s、實際 2–3 s），協定已將 dwell 顯式改為 3 s 以保證板端 1 Hz 取樣"
                 "可得穩定樣本。板端域係數之線上複測列為 [待測]（Round 41）。")
    insert_after(pb, "三方法（含 EKF）之整體精度將於 4.4 節在同一框架下比較。")

    # ============ 4.4 ============
    set_p(P("本節為三法之公平比較。比較框架"),
          "本節為三法之公平比較，數值除特別註記者外均為 Round 40 整輪板端**實測**"
          "（同電池、同協定、同 MCU、同量測路徑；第三章 + 4.0 節）。")

    set_p(P("於同一組放電 cycle 上"),
          "於同一輪四個放電 cycle（0.5／1.0／1.5／2.0C）上，以 4.0 節框架對三法計算精度指標"
          "（各法取四倍率之範圍）：")

    tb4 = T("方法", "RMSE", "備註")
    data4 = [["庫倫計數（板端）", "1.8–2.2", "1.5–1.9", "3.2–3.8", "N/A（無收斂機制）",
              "恆負偏；來源為板端與台架兩條電流量測鏈之刻度差（約 +1.2%），非演算法誤差"],
             ["EKF（板端）", "0.87–4.5", "0.85–4.2", "1.1–6.4", "單步（電壓播種）",
              "誤差隨倍率增大、恆正偏（R_0 量測域錯位，見 4.2.5）；0.5C 時 0.87% 為三法最佳"],
             ["動態阻抗（單獨反推，離線）", "6–10", "5–8", "~20", "首事件（≤60 s）",
              "fresh rounds 1–3 實測（4.3.4）；中段誤差較大"],
             ["動態阻抗（線上，台架域係數）", "~28", "—", "~50", "—",
              "量測域錯位之失效案例（4.3.4）；板端域係數複測 [待測]"]]
    while len(tb4.rows) < 1 + len(data4):
        tb4._tbl.append(copy.deepcopy(tb4.rows[-1]._tr))
    set_cell(tb4.rows[0].cells[4], "收斂時間")
    for i, row in enumerate(data4, start=1):
        for j, v in enumerate(row):
            set_cell(tb4.rows[i].cells[j], v)

    p410 = P("表 4-4　三法標稱工況精度比較")
    set_p(p410, "表 4-4　三法標稱工況精度比較（Round 40 板端實測；動態阻抗離線列為 "
                "rounds 1–3 實測）")
    insert_after(p410,
                 "表中有兩個「誤差來源不在演算法本身」的實測教訓值得強調：庫倫計數的負偏來自"
                 "**電流量測鏈刻度**、EKF 的正偏與動態阻抗的失效均來自**參數辨識域與部署量測域"
                 "不一致**。三者共同指向：嵌入式 SOC 估測之精度上限，經常由量測鏈一致性而非演算法"
                 "複雜度決定。")

    set_p(P("三項壓力測試（皆 [待測]"),
          "三項壓力測試（初值錯誤恢復已於板端實測；後二項 [待測]）：")
    set_p(P("1. 初值錯誤恢復：初值偏移"),
          "1. **初值錯誤恢復（實測）**：庫倫計數無修正機制，錯誤初值之偏差永久保留——本研究於 "
          "Round 40 前段實際觀察到此現象（開機錨定於非滿電狀態時，其誤差直到下一次充飽重錨才消除），"
          "並據此於韌體加入「充飽自動重錨」機制（V > 4.15 V 且 |I| < 20 mA 持續 60 s，整輪 "
          "4/4 次正確觸發）。EKF 以除錯介面強制偏移 +25% 後，一次量測更新即拉回（4.2.5）。"
          "動態阻抗於首個擾動事件（本協定 ≤60 s）即獨立重估，實測每事件均重新錨定。")
    set_p(P("2. C-rate 切換：於放電中切換倍率"),
          "2. **C-rate 切換**：於放電中切換倍率，觀察誤差尖峰。預期 EKF 因 RC 模型可吸收暫態、"
          "動態阻抗於切換瞬間 Δ I 異常需剔除該事件（韌體已設 |Δ I| 上限與放電向過濾）。[待測]。")
    set_p(P("3. 噪聲抗性：人為注入"),
          "3. **噪聲抗性**：人為注入電流／電壓量測噪聲，比較三法 SOC 抖動。預期 EKF 因 Q/R "
          "濾波最穩、庫倫計數對電流噪聲積分平均、動態阻抗對單點差分噪聲最敏感。[待測]。")

    tb5 = T("壓力測試", "庫倫計數", "動態阻抗")
    set_cell(tb5.rows[1].cells[0], "初值錯誤恢復")
    set_cell(tb5.rows[1].cells[1], "不收斂（實測；以充飽重錨補救）")
    set_cell(tb5.rows[1].cells[2], "單步拉回（實測）")
    set_cell(tb5.rows[1].cells[3], "首事件重估 ≤60 s（實測）")
    set_p(P("表 4-5　三法強健性比較"),
          "表 4-5　三法強健性比較（初值列為 Round 40 板端實測）")

    tb6 = T("方法", "程式空間", "是否需浮點")
    hdr6 = ["方法", "程式空間 (B)", "記憶體 (B)",
            "每次更新運算量（cycles 中位，@64 MHz）", "是否需浮點", "備註"]
    for j, h in enumerate(hdr6):
        set_cell(tb6.rows[0].cells[j], h)
    data6 = [["庫倫計數", "+1,124", "+20", "305（4.8 µs）", "否", "資源下界（int64 累加）"],
             ["EKF", "+1,868", "+40", "15,187（237 µs）", "是（軟浮點）",
              "含 21 點開路電壓對照表（168 B）"],
             ["動態阻抗（無事件）", "+1,456", "+48", "432（6.8 µs）", "是（軟浮點）",
              "庫倫內插路徑"],
             ["動態阻抗（有事件）", "同上", "同上", "3,117（48.7 µs）", "是（軟浮點）",
              "差分＋開根＋二次反解，每 60 s 一次"]]
    while len(tb6.rows) < 1 + len(data6):
        tb6._tbl.append(copy.deepcopy(tb6.rows[-1]._tr))
    for i, row in enumerate(data6, start=1):
        for j, v in enumerate(row):
            set_cell(tb6.rows[i].cells[j], v)

    set_p(P("表 4-6　三法嵌入式資源佔用比較"),
          "表 4-6　三法嵌入式資源佔用實測（STM32G071 @64 MHz、-Og、軟浮點；程式/記憶體為對"
          "共用骨架之差分量測，運算量為整輪 70,802 次更新之統計中位數）。EKF 之每次更新運算量"
          "約為庫倫計數之 50 倍，但摺算後仍僅佔 1 s 更新週期之 0.024%——三法在本平台皆遠未逼近"
          "即時性極限，資源差異之意義在更低階 MCU 或更高更新率場景。")

    set_p(P("此表預期清楚呈現「精度—資源」權衡"),
          "此表清楚呈現「精度—資源」權衡：庫倫計數最省資源但無自我修正能力；EKF 精度與強健性最佳，"
          "程式空間、記憶體與運算量代價亦最高；動態阻抗居中且實驗成本最低。此即第二章 2.3 節所述"
          "「商用晶片為何多捨 EKF」之量化佐證——本研究以同硬體實測填補文獻少見之缺口，且實測顯示"
          "該取捨在 Cortex-M0+ 級距上已不構成即時性障礙（見表 4-6 註）。")

    set_p(P("綜合精度、強健性與資源三軸"),
          "綜合精度、強健性與資源三軸之實測結果（4.4.1–4.4.3），依目標應用之約束給出方法選用建議：")

    tb7 = T("應用約束", "建議方法")
    set_cell(tb7.rows[0].cells[2], "理由（實測佐證）")
    set_cell(tb7.rows[1].cells[1], "庫倫計數 + 充飽自動重錨")
    set_cell(tb7.rows[1].cells[2],
             "資源下界（+1.1 KB／305 cycles）；重錨機制實測可消除初值偏移，"
             "精度上限由電流量測鏈刻度決定（±2%）")
    set_cell(tb7.rows[2].cells[2],
             "低倍率精度最佳（0.87%）且單步自錯誤初值收斂；資源代價 +1.9 KB／237 µs 在 "
             "Cortex-M0+ 上仍僅佔更新週期 0.024%；參數辨識須與部署量測鏈同域")
    set_cell(tb7.rows[3].cells[1], "動態阻抗（＋庫倫內插）")
    set_cell(tb7.rows[3].cells[2],
             "無須初值與靜置、首事件 ≤60 s 重估；但反推精度有限（6–10%）且對照表對量測域高度敏感")

    # ============ 4.5 ============
    p428 = P("本章於第三章建立的統一平台上")
    set_p(p428,
          "本章於第三章建立的統一平台上，依「原理 → 演算法 → 嵌入式實作 → 實驗結果」一致結構，"
          "實作並比較三種 SOC 估測方法，三法均已完成嵌入式移植並以同一電池、同一協定、同一微控制器"
          "整輪實測（Round 40）。實測結論可濃縮為三點：")
    q1 = insert_after(p428,
          "其一，**精度排序依工況而異**：低倍率下 EKF 最佳（0.5C 時 RMSE 0.87%），庫倫計數穩定"
          "居中（1.8–2.2%，且其負偏可溯源至兩條電流量測鏈之刻度差而非演算法），動態阻抗單獨反推"
          "最弱（6–10%、中段病態）。")
    q2 = insert_after(q1,
          "其二，**資源階梯已被量化**（表 4-6）：庫倫計數 305 cycles／EKF 15,187 cycles／"
          "動態阻抗 432（無事件）與 3,117（有事件）cycles，EKF 之運算代價約為庫倫之 50 倍——"
          "但於 64 MHz Cortex-M0+ 上折算僅 237 µs，佔 1 s 更新週期不到萬分之三。「EKF 太重」"
          "在本級距 MCU 上並不成立，資源論述的分水嶺在更低階平台。")
    insert_after(q2,
          "其三，也是本章最重要的工程發現：**量測域一致性凌駕演算法選擇**。EKF 的倍率相依正偏"
          "（R_0 辨識域含線材電阻）與動態阻抗線上失效（台架域對照表最低點高於板端一切量測值）為"
          "同一根因的兩個表現——參數與對照表若非與部署端量測鏈同域建立，演算法設計再正確也會"
          "系統性失準。此實測教訓直接回應第一章 1.2 節「評估環境失真」之研究缺口，亦是唯有在真實"
          "嵌入式平台上做公平比較才可能暴露的問題。")
    set_p(P("凡尚未實測之數值均誠實標註"),
          "尚未完成者（表 4-5 之 C-rate 切換與噪聲抗性壓力測試、量測域修正後之 Round 41 複測）"
          "均誠實標註 [待測]。第五章將在此基礎上，討論三法並行於同一微控制器的整合策略、運算負載"
          "分配、fresh-cell 三輪變異性，以及實驗過程之工程教訓。")

    d.save(DST)
    print(f"[done] {DST}")


if __name__ == "__main__":
    main()
